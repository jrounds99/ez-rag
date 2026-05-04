"""Cross-platform downloader for the public-domain RAG sample corpus.

Reads `sample_data/curation.json` and downloads every entry to
`sample_data/fetched/<save_as>`. Idempotent — re-running skips files
that already exist. Writes `sample_data/manifest.json` recording
what was actually fetched, file sizes, SHA-256 hashes, and any
URLs that 404'd.

Usage:
    python sample_data/fetch.py                     # full default fetch
    python sample_data/fetch.py --topic geology     # only geology entries
    python sample_data/fetch.py --max-mb 30         # cap total bytes
    python sample_data/fetch.py --dry-run           # list, don't download
    python sample_data/fetch.py --no-convert        # skip docx conversion
    python sample_data/fetch.py --force             # re-download even if present

After fetch, the corpus lives at `sample_data/fetched/`. Point
ez-rag at it:

    python -m ez_rag.cli ingest --workspace sample_data/fetched

…or hand it to the bench:

    ./bench/run.sh search --workspace sample_data/fetched
"""
from __future__ import annotations

import argparse
import hashlib
import io
import json
import shutil
import subprocess
import sys
import time
import zipfile
from pathlib import Path
from typing import Optional


HERE = Path(__file__).resolve().parent
CURATION = HERE / "curation.json"
FETCHED = HERE / "fetched"
MANIFEST = HERE / "manifest.json"


# ============================================================================
# IO
# ============================================================================

def _say(msg: str) -> None:
    print(msg, flush=True)


def _human_bytes(n: int) -> str:
    n = float(n)
    for unit in ("B", "KB", "MB", "GB"):
        if n < 1024 or unit == "GB":
            return f"{n:.1f} {unit}" if unit != "B" else f"{int(n)} B"
        n /= 1024
    return f"{n:.1f} GB"


def _sha256_of(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


# ============================================================================
# HTTP — tries httpx first (richer error reporting), falls back to urllib
# ============================================================================

def _download(url: str, dest: Path, *, timeout: float = 60.0) -> tuple[bool, str]:
    """Download `url` to `dest`. Returns (success, message_or_size).
    Streams to disk in 1 MB chunks. Honors HTTP redirects."""
    dest.parent.mkdir(parents=True, exist_ok=True)
    tmp = dest.with_suffix(dest.suffix + ".part")
    try:
        try:
            import httpx  # type: ignore
            with httpx.stream(
                "GET", url, timeout=timeout, follow_redirects=True,
                headers={"User-Agent": "ez-rag-sample-fetch/1.0"},
            ) as r:
                if r.status_code != 200:
                    return False, f"HTTP {r.status_code}"
                with tmp.open("wb") as f:
                    for chunk in r.iter_bytes(chunk_size=1024 * 1024):
                        f.write(chunk)
        except ImportError:
            # Standard-lib fallback so the script works on a fresh
            # Python install without `pip install httpx`.
            import urllib.request
            req = urllib.request.Request(
                url,
                headers={"User-Agent": "ez-rag-sample-fetch/1.0"},
            )
            with urllib.request.urlopen(req, timeout=timeout) as resp:
                if resp.status != 200:
                    return False, f"HTTP {resp.status}"
                with tmp.open("wb") as f:
                    shutil.copyfileobj(resp, f)
        tmp.replace(dest)
        return True, _human_bytes(dest.stat().st_size)
    except Exception as ex:
        if tmp.exists():
            try:
                tmp.unlink()
            except OSError:
                pass
        return False, f"{type(ex).__name__}: {ex}"


# ============================================================================
# Format conversion (PDF -> DOCX)
# ============================================================================

def _have_pandoc() -> bool:
    return shutil.which("pandoc") is not None


def _convert_pdf_to_docx(pdf_path: Path, docx_path: Path) -> tuple[bool, str]:
    """Best-effort PDF → DOCX conversion. Tries pandoc first, then
    a python-docx fallback that wraps extracted PDF text in a Word
    document. Returns (ok, message)."""
    docx_path.parent.mkdir(parents=True, exist_ok=True)

    # Pandoc path — preserves a lot more structure but requires the
    # binary AND a working LaTeX/typst path for some PDFs. We retry
    # with --extract-media if the first run fails.
    if _have_pandoc():
        try:
            r = subprocess.run(
                ["pandoc", str(pdf_path), "-o", str(docx_path)],
                capture_output=True, text=True, timeout=120,
            )
            if r.returncode == 0 and docx_path.is_file():
                return True, "pandoc"
        except (subprocess.TimeoutExpired, OSError):
            pass

    # Fallback — extract text via pypdf and write a plain DOCX with
    # python-docx. Loses formatting but produces a valid .docx file
    # that ez-rag's parser can ingest, exercising the docx code path.
    try:
        from pypdf import PdfReader  # type: ignore
        from docx import Document  # type: ignore
    except ImportError as ex:
        return False, f"missing dep ({ex.name})"
    try:
        reader = PdfReader(str(pdf_path))
        doc = Document()
        doc.add_heading(pdf_path.stem.replace("-", " ").title(), 0)
        for i, page in enumerate(reader.pages, start=1):
            try:
                txt = page.extract_text() or ""
            except Exception:
                txt = ""
            txt = txt.strip()
            if not txt:
                continue
            doc.add_heading(f"Page {i}", level=2)
            for para in txt.split("\n\n"):
                p = para.strip()
                if p:
                    doc.add_paragraph(p)
        doc.save(str(docx_path))
        return True, "python-docx fallback"
    except Exception as ex:
        return False, f"{type(ex).__name__}: {ex}"


# ============================================================================
# ZIP extraction (for the BLS QCEW package)
# ============================================================================

def _extract_zip(zip_path: Path, target_dir: Path) -> list[Path]:
    """Extract a zip into a sibling directory named after the zip
    (without .zip). Returns list of extracted files. Skips already-
    extracted output."""
    target_dir.mkdir(parents=True, exist_ok=True)
    extracted: list[Path] = []
    try:
        with zipfile.ZipFile(zip_path) as z:
            for name in z.namelist():
                # Defensive — reject any path traversal attempt
                clean = Path(name).as_posix().lstrip("/")
                if ".." in clean.split("/"):
                    continue
                out = target_dir / clean
                if out.is_file() and out.stat().st_size > 0:
                    extracted.append(out)
                    continue
                out.parent.mkdir(parents=True, exist_ok=True)
                with z.open(name) as src, out.open("wb") as dst:
                    shutil.copyfileobj(src, dst)
                extracted.append(out)
    except zipfile.BadZipFile:
        return []
    return extracted


# ============================================================================
# Main flow
# ============================================================================

def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--topic", default=None,
                     help="Only fetch entries with this topic (e.g. geology).")
    ap.add_argument("--max-mb", type=float, default=0.0,
                     help="Stop after fetching this many MB (0 = no cap).")
    ap.add_argument("--dry-run", action="store_true",
                     help="List what would be fetched, don't download.")
    ap.add_argument("--no-convert", action="store_true",
                     help="Skip the post-processing PDF -> DOCX step.")
    ap.add_argument("--force", action="store_true",
                     help="Re-download files that already exist.")
    ap.add_argument("--skip-optional", action="store_true",
                     default=True,
                     help="Skip entries marked 'optional: true' in "
                          "curation.json (large datasets). Default: True. "
                          "Pass --include-optional to fetch them.")
    ap.add_argument("--include-optional", action="store_true",
                     help="Include large optional datasets (e.g. BLS QCEW).")
    args = ap.parse_args()
    if args.include_optional:
        args.skip_optional = False

    if not CURATION.is_file():
        _say(f"[!] curation.json not found at {CURATION}")
        return 1

    curation = json.loads(CURATION.read_text(encoding="utf-8"))
    items = curation.get("items", [])
    post = curation.get("post_processing", {})

    if args.topic:
        items = [it for it in items if it.get("topic") == args.topic]
        _say(f"[topic={args.topic}] {len(items)} item(s) match")
    if args.skip_optional:
        before = len(items)
        items = [it for it in items if not it.get("optional", False)]
        if before != len(items):
            _say(f"  skipping {before - len(items)} optional "
                 f"(large) entr(y/ies). Pass --include-optional to fetch.")

    FETCHED.mkdir(parents=True, exist_ok=True)

    started = time.time()
    fetched: list[dict] = []
    skipped: list[dict] = []
    failed: list[dict] = []
    bytes_so_far = 0.0

    for i, it in enumerate(items, start=1):
        url = it["url"]
        save_as = it["save_as"]
        dest = FETCHED / save_as
        size_mb = float(it.get("expected_size_mb", 0))

        if args.max_mb and (bytes_so_far / (1024 * 1024)) >= args.max_mb:
            _say(f"  [{i}/{len(items)}] STOP: --max-mb {args.max_mb} reached")
            break

        if args.dry_run:
            _say(f"  [{i}/{len(items)}] would fetch  "
                 f"({size_mb:>5.1f} MB)  {save_as}")
            continue

        if dest.is_file() and dest.stat().st_size > 0 and not args.force:
            _say(f"  [{i}/{len(items)}] skip (exists) "
                 f"{_human_bytes(dest.stat().st_size):>9}  {save_as}")
            skipped.append({"item": save_as, "size_bytes": dest.stat().st_size})
            bytes_so_far += dest.stat().st_size
            continue

        _say(f"  [{i}/{len(items)}] fetching     "
             f"({size_mb:>5.1f} MB)  {save_as}")
        ok, msg = _download(url, dest)
        if not ok:
            _say(f"               × FAILED: {msg}")
            failed.append({"item": save_as, "url": url, "error": msg})
            continue
        actual = dest.stat().st_size
        bytes_so_far += actual
        fetched.append({
            "item": save_as,
            "url": url,
            "size_bytes": actual,
            "sha256": _sha256_of(dest),
            "title": it.get("title", ""),
            "format": it.get("format", ""),
            "topic": it.get("topic", ""),
            "states": it.get("states", []),
            "license": it.get("license", ""),
        })
        _say(f"               OK {msg}")

        # Auto-extract zip archives if the curation entry asks for it
        if it.get("extract_after") and dest.suffix.lower() == ".zip":
            target = dest.with_suffix("")
            extracted = _extract_zip(dest, target)
            _say(f"               extracted {len(extracted)} "
                 f"file(s) -> {target.relative_to(FETCHED)}")

    # Post-processing — PDF -> DOCX conversions for format diversity
    converted: list[dict] = []
    convert_failed: list[dict] = []
    if not args.dry_run and not args.no_convert:
        for entry in post.get("convert_to_docx", []):
            src = FETCHED / entry["source"]
            tgt = FETCHED / entry["target"]
            if not src.is_file():
                continue
            if tgt.is_file() and not args.force:
                _say(f"  convert skip (exists) {entry['target']}")
                continue
            _say(f"  converting {entry['source']} -> {entry['target']}")
            ok, msg = _convert_pdf_to_docx(src, tgt)
            if ok:
                converted.append({
                    "source": entry["source"],
                    "target": entry["target"],
                    "method": msg,
                    "size_bytes": tgt.stat().st_size,
                })
                _say(f"               OK ({msg})")
            else:
                convert_failed.append({
                    "source": entry["source"],
                    "target": entry["target"],
                    "error": msg,
                })
                _say(f"               × {msg}")

    # Manifest
    duration_s = time.time() - started
    manifest = {
        "fetched_at": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()),
        "duration_s": round(duration_s, 1),
        "total_size_bytes": int(bytes_so_far),
        "total_size_human": _human_bytes(int(bytes_so_far)),
        "fetched_count": len(fetched),
        "skipped_count": len(skipped),
        "failed_count": len(failed),
        "fetched": fetched,
        "skipped": skipped,
        "failed": failed,
        "converted": converted,
        "convert_failed": convert_failed,
    }
    if not args.dry_run:
        MANIFEST.write_text(
            json.dumps(manifest, indent=2),
            encoding="utf-8",
        )

    _say("")
    _say(f"=== sample_data fetch summary ===")
    _say(f"  fetched:   {len(fetched)}")
    _say(f"  skipped:   {len(skipped)} (already on disk)")
    _say(f"  failed:    {len(failed)}")
    _say(f"  converted: {len(converted)} (PDF -> DOCX)")
    _say(f"  total:     {_human_bytes(int(bytes_so_far))}")
    if not args.dry_run:
        _say(f"  manifest:  {MANIFEST}")
    if failed:
        _say("")
        _say("Some downloads failed. Common causes:")
        _say("  - Government URL was reorganized (404)")
        _say("  - Temporary network / TLS error")
        _say("Re-run to retry; or open curation.json and update the URL.")
        for f in failed:
            _say(f"  • {f['item']}: {f['error']}")

    return 0 if not failed else 2


if __name__ == "__main__":
    sys.exit(main())
