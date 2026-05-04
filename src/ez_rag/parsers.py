"""Document parsers. Each returns a list[ParsedSection] which the chunker consumes.

A ParsedSection is text plus optional location metadata (page, section path).
We deliberately keep parsers in pure Python with light deps.
"""
from __future__ import annotations

import csv
import email
import re
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Callable

# ----- types -----------------------------------------------------------------

@dataclass
class ParsedSection:
    text: str
    page: int | None = None
    section: str = ""
    meta: dict = field(default_factory=dict)


# ----- registry --------------------------------------------------------------

ParserFn = Callable[[Path], list[ParsedSection]]
_REGISTRY: dict[str, ParserFn] = {}

def register(*exts: str):
    def deco(fn: ParserFn) -> ParserFn:
        for e in exts:
            _REGISTRY[e.lower()] = fn
        return fn
    return deco

def get_parser(path: Path) -> ParserFn | None:
    return _REGISTRY.get(path.suffix.lower())

def supported_extensions() -> set[str]:
    return set(_REGISTRY.keys())


# ----- helpers ---------------------------------------------------------------

def _collapse_table_runs(text: str) -> str:
    """Collapse pathological table-column-flattening runs.

    PDFs are 2D, but text extraction serializes them to 1D. When a table
    has many rows with the same short value in one column (e.g. a Spells
    table where most spells have Ritual=No), extraction often produces:

        Ritual
        No
        No
        No
        No
        ...   (× 47 more)

    Without context (spell name / school / level), each "No" is
    meaningless — but it competes for chunk space and pollutes
    retrieval. Detect runs of 6+ identical short (≤ 12 char) lines and
    collapse to "<value> (×N)" so the table column is acknowledged
    once but doesn't dominate.

    Conservative thresholds — won't touch normal repetition like a
    poem refrain or a few sequential bullet points.
    """
    if not text:
        return text
    lines = text.split("\n")
    out: list[str] = []
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()
        # Only consider short non-empty lines as "column-cell candidates"
        if stripped and len(stripped) <= 12:
            j = i + 1
            while j < len(lines) and lines[j].strip() == stripped:
                j += 1
            run = j - i
            if run >= 6:
                # Preserve any leading whitespace from the first line so
                # bulleted/indented content still looks consistent.
                lead = lines[i][: len(lines[i]) - len(lines[i].lstrip())]
                out.append(f"{lead}{stripped} (×{run})")
                i = j
                continue
        out.append(lines[i])
        i += 1
    return "\n".join(out)


def _normalize(text: str) -> str:
    text = text.replace("\x00", "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = _collapse_table_runs(text)
    return text.strip()


# ----- TXT / MD / RST --------------------------------------------------------

@register(".txt", ".md", ".markdown", ".rst", ".log")
def parse_text(path: Path) -> list[ParsedSection]:
    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        text = path.read_text(encoding="latin-1", errors="replace")
    return [ParsedSection(text=_normalize(text))]


# ----- PDF -------------------------------------------------------------------

# Heuristic thresholds for "this page's text is garbled" detection. Tuned
# against examples like:
#   "hAppe�d to the \pell\ Mor&e�kAi�e�'l Bo1A�tif1AI"
#   "ewAr&'\ I/qt Air, A�& All the re\t ?"
# which are typical when a PDF embeds a custom/subsetted font with a
# broken or missing ToUnicode cmap — pypdf falls back to glyph IDs and
# you get nonsense.
_GARBLED_REPLACEMENT_RATIO = 0.02   # >2% replacement chars (U+FFFD) → garbled
_GARBLED_BACKSLASH_RATIO   = 0.025  # >2.5% backslashes → garbled escape sequences
_GARBLED_VOWEL_FLOOR       = 0.20   # English text is ~35-40% vowels of alpha chars
_GARBLED_NONALNUM_CEIL     = 0.45   # >45% non-alphanumeric (excl. whitespace) → garbled


def _looks_like_toc_fragment(text: str) -> bool:
    """Heuristic: text is a Table-of-Contents-style index page that
    extracted as fragmentary "label / page-number" rows.

    Even when font extraction works, TOC pages are search-poison: a
    chunk that says "Fighter\\n59\\nMonk\\n.61\\nOmin Dran.\\n..196"
    will get retrieved for queries like "fighter abilities" and
    waste a top-K slot on a page-number index entry.

    OCR pages look like this AFTER recovery sometimes — the user
    flagged this in a screenshot. We drop these.

    Signals:
      - many short lines (< 25 chars)
      - many lines that are just numbers (or "..196" / ".61"
        leftover dot-leader fragments)
      - average line length << prose
    """
    if not text:
        return False
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    if len(lines) < 6:
        return False  # too short to confidently classify

    short = sum(1 for l in lines if len(l) <= 25)
    # Lines that are essentially a page number — possibly with leading
    # dots from collapsed dot-leaders ("..196", ".61").
    page_num_like = sum(
        1 for l in lines
        if l.lstrip(".").strip().isdigit() and len(l.lstrip(".").strip()) <= 4
    )
    avg_line_len = sum(len(l) for l in lines) / len(lines)

    # Strong TOC signal: >40% short lines AND >15% bare page numbers
    # AND short average line length. Tuned to match the user's example
    # without flagging normal short-paragraph prose or bullet lists.
    if (short / len(lines) > 0.40
            and page_num_like / len(lines) > 0.15
            and avg_line_len < 22):
        return True
    return False


def _text_looks_garbled(text: str) -> bool:
    """Heuristic: does this page's extracted text look like font-cmap
    garbage rather than real prose?

    Returns True for short-vowel / high-replacement / high-backslash /
    high-symbol-ratio text. False for short snippets (< 40 chars), so
    headers like "Chapter 1" don't trip the detector.
    """
    if not text:
        return False
    t = text.strip()
    if len(t) < 40:
        return False

    n = len(t)
    # 1. Replacement characters (the � you saw)
    n_replacement = t.count("�")
    if n / max(1, n_replacement) and n_replacement / n > _GARBLED_REPLACEMENT_RATIO:
        return True
    # 2. Backslash escape ratio — broken cmaps emit \pell\ \td\ etc.
    n_back = t.count("\\")
    if n_back / n > _GARBLED_BACKSLASH_RATIO:
        return True
    # 3. Vowel ratio in alpha chars. English prose is ~35-40%; if a page's
    # alpha chars are <20% vowels something's badly wrong.
    alpha = [c for c in t.lower() if c.isalpha()]
    if alpha and len(alpha) > 50:
        vowels = sum(1 for c in alpha if c in "aeiou")
        if vowels / len(alpha) < _GARBLED_VOWEL_FLOOR:
            return True
    # 4. Non-alphanumeric / non-whitespace ratio — broken extraction often
    # produces dense punctuation soup like "Mor&e�kAi�e�'l Bo1A�tif1AI".
    non_alnum_ws = sum(1 for c in t if not c.isalnum() and not c.isspace())
    if non_alnum_ws / n > _GARBLED_NONALNUM_CEIL:
        return True
    return False


@register(".pdf")
def parse_pdf(path: Path, on_progress=None, on_recovery=None,
              permissive: bool = False) -> list[ParsedSection]:
    """Parse a PDF page-by-page.

    `on_progress(page, total, ocr=False)` — if given, called after each page
    so callers can surface live progress for big PDFs (which can otherwise
    block for minutes with no UI feedback).

    Garbled-page recovery: pypdf's `extract_text()` produces nonsense when
    a PDF uses a custom font with a broken/missing ToUnicode cmap. We
    detect this per-page via `_text_looks_garbled()` and re-extract via
    OCR for those pages only — keeps fast pypdf extraction for clean
    pages, falls back to OCR (which reads pixels, not fonts) for the
    broken ones.

    `permissive` — when True, OCR results that still look garbled or look
    like TOC fragments are KEPT (with meta.questionable=True) instead of
    being dropped. Caller (typically ingest.py with llm_correct_garbled
    enabled) is expected to feed those sections through an LLM cleanup
    pass and drop them only if the LLM also fails. Default False
    preserves the original "drop poison upstream" behavior.
    """
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError as e:
        raise RuntimeError("pypdf not installed; required for PDF parsing") from e
    sections: list[ParsedSection] = []
    reader = PdfReader(str(path))
    total = len(reader.pages)
    page_texts: list[str] = []
    garbled_pages: list[int] = []      # 1-indexed page numbers
    for i, page in enumerate(reader.pages, start=1):
        try:
            t = page.extract_text() or ""
        except Exception:
            t = ""
        page_texts.append(t)
        if t.strip():
            if _text_looks_garbled(t):
                garbled_pages.append(i)
                # Don't add the garbled section yet — we'll try to OCR
                # it below. Add a placeholder so indexing stays stable.
                sections.append(ParsedSection(
                    text="", page=i,
                    meta={"reparse_pending": "garbled"},
                ))
            else:
                sections.append(ParsedSection(text=_normalize(t), page=i))
        if on_progress:
            try:
                on_progress(i, total, ocr=False)
            except Exception:
                pass

    # ----- Recovery passes -----
    total_chars = sum(len(t) for t in page_texts)

    # 1. Whole-document fallback when pypdf got almost nothing — the PDF
    #    is probably a scan with no text layer. OCR every page.
    if total_chars < 50 * max(1, total):
        ocr_sections = _ocr_pdf_pages(path, on_progress=on_progress)
        if ocr_sections:
            return ocr_sections

    # 2. Per-page fallback: pypdf got SOMETHING but specific pages are
    #    garbled. OCR just those pages and substitute their text in.
    if garbled_pages:
        # Capture the BEFORE text so the GUI can show side-by-side
        # before/after when previewing a recovery.
        before_map = {p: page_texts[p - 1] for p in garbled_pages}
        ocr_map = _ocr_pdf_pages_subset(
            path, garbled_pages,
            on_progress=on_progress,
            on_recovery=on_recovery,
            before_texts=before_map,
        )
        # Replace the placeholder ParsedSection with the OCR text where
        # available; drop the placeholder if OCR also returned nothing.
        rebuilt: list[ParsedSection] = []
        for sec in sections:
            if sec.meta.get("reparse_pending") == "garbled":
                ocr_text = ocr_map.get(sec.page or 0, "").strip()
                # Three rejection criteria — we drop the page if ANY hit:
                #  1. OCR returned nothing
                #  2. OCR returned text that's STILL garbled (different
                #     glyphs from a sub-image, weird unicode, etc.)
                #  3. OCR returned a TOC-fragment index page that's
                #     fragmentary "label / page-number" garbage with
                #     no retrieval value (the user's "Preface / Fighter
                #     / 59 / Monk / .61" example)
                still_bad = (
                    not ocr_text
                    or _text_looks_garbled(ocr_text)
                    or _looks_like_toc_fragment(ocr_text)
                )
                if still_bad and not permissive:
                    # Drop the page rather than poison the index.
                    pass
                elif still_bad and permissive and ocr_text:
                    # Keep it but flag as questionable so the caller's
                    # LLM-correction pass can take a shot. If the LLM
                    # also rejects, the caller drops it then.
                    rebuilt.append(ParsedSection(
                        text=_normalize(ocr_text),
                        page=sec.page,
                        meta={"ocr": True, "reparse": "garbled",
                              "questionable": True},
                    ))
                elif not still_bad:
                    rebuilt.append(ParsedSection(
                        text=_normalize(ocr_text),
                        page=sec.page,
                        meta={"ocr": True, "reparse": "garbled"},
                    ))
            else:
                rebuilt.append(sec)
        sections = rebuilt
    return sections


def _ocr_pdf_pages_subset(
    path: Path, pages_1indexed: list[int],
    *,
    on_progress=None,
    on_recovery=None,
    before_texts: dict[int, str] | None = None,
) -> dict[int, str]:
    """Render specific PDF pages to images and OCR them. Returns
    {page_number: text}. Used to recover individual pages where pypdf's
    extraction looked garbled.

    `on_recovery(payload)` — optional. When given, fires once per page
    after both the page render and OCR have finished, with a dict shape:
        {file, page, image_path, before, after}
    The GUI uses this to show a live before/after preview during ingest.
    Saving the rendered image to the preview cache costs ~50 ms + ~200 KB
    per page; only do it when the caller actually wants the preview.
    """
    out: dict[int, str] = {}
    if not pages_1indexed:
        return out
    try:
        import pypdfium2 as pdfium  # type: ignore
    except ImportError:
        return out
    from .ocr import ocr_image
    from .preview import cache_path_for as _img_cache_path
    try:
        pdf = pdfium.PdfDocument(str(path))
    except Exception:
        return out
    n_pages = len(pdf)
    pages_1indexed = [p for p in pages_1indexed if 1 <= p <= n_pages]
    for j, p in enumerate(pages_1indexed, start=1):
        text = ""
        img_path = ""
        try:
            page = pdf[p - 1]
            bitmap = page.render(scale=2.0)
            pil = bitmap.to_pil()
            # Save the rendered page to the preview cache only if a
            # recovery callback wants to display it. Skips the disk
            # write entirely when previews are off.
            if on_recovery:
                try:
                    cache_p = _img_cache_path(path, p)
                    pil.save(cache_p, format="PNG", optimize=True)
                    img_path = str(cache_p)
                except Exception:
                    img_path = ""
            text = ocr_image(pil) or ""
            out[p] = text
        except Exception:
            out[p] = ""

        if on_recovery:
            try:
                before = (before_texts or {}).get(p, "")
                on_recovery({
                    "file": str(path),
                    "page": p,
                    "image_path": img_path,
                    "before": (before or "")[:1500],
                    "after": (text or "")[:1500],
                })
            except Exception:
                pass

        if on_progress:
            try:
                on_progress(j, len(pages_1indexed), ocr=True)
            except Exception:
                pass
    return out


def _ocr_pdf_pages(path: Path, on_progress=None) -> list[ParsedSection]:
    """Render PDF pages to images and OCR them. Returns [] if deps missing."""
    try:
        from pypdf import PdfReader  # noqa: F401  (already required)
    except ImportError:
        return []
    # Try pypdfium2 (Apache + BSD) first; it's the lightest renderer.
    try:
        import pypdfium2 as pdfium  # type: ignore
    except ImportError:
        return []
    from .ocr import ocr_image  # local import to avoid circular at top
    sections: list[ParsedSection] = []
    pdf = pdfium.PdfDocument(str(path))
    total = len(pdf)
    for i in range(total):
        page = pdf[i]
        bitmap = page.render(scale=2.0)
        pil = bitmap.to_pil()
        text = ocr_image(pil) or ""
        if text.strip():
            sections.append(ParsedSection(text=_normalize(text), page=i + 1, meta={"ocr": True}))
        if on_progress:
            try:
                on_progress(i + 1, total, ocr=True)
            except Exception:
                pass
    return sections


# ----- DOCX ------------------------------------------------------------------

@register(".docx")
def parse_docx(path: Path) -> list[ParsedSection]:
    try:
        import docx  # type: ignore  (python-docx)
    except ImportError as e:
        raise RuntimeError("python-docx not installed; required for DOCX") from e
    doc = docx.Document(str(path))
    sections: list[ParsedSection] = []
    current_heading = ""
    buf: list[str] = []
    def flush():
        if buf:
            sections.append(ParsedSection(text=_normalize("\n".join(buf)),
                                          section=current_heading))
            buf.clear()
    for para in doc.paragraphs:
        style = (para.style.name or "").lower() if para.style else ""
        if style.startswith("heading"):
            flush()
            current_heading = para.text.strip()
            if current_heading:
                buf.append(f"# {current_heading}")
        else:
            if para.text.strip():
                buf.append(para.text)
    # Tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append(" | ".join(cell.text.strip() for cell in row.cells))
        if rows:
            buf.append("\n".join(rows))
    flush()
    return sections


# ----- XLSX / CSV ------------------------------------------------------------

@register(".xlsx", ".xlsm")
def parse_xlsx(path: Path) -> list[ParsedSection]:
    try:
        import openpyxl  # type: ignore
    except ImportError as e:
        raise RuntimeError("openpyxl not installed; required for XLSX") from e
    wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    sections: list[ParsedSection] = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = ["" if v is None else str(v) for v in row]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            sections.append(ParsedSection(
                text=_normalize("\n".join(rows)),
                section=sheet_name,
            ))
    return sections


@register(".csv", ".tsv")
def parse_csv(path: Path) -> list[ParsedSection]:
    delim = "\t" if path.suffix.lower() == ".tsv" else ","
    rows: list[str] = []
    with path.open("r", encoding="utf-8", errors="replace", newline="") as f:
        reader = csv.reader(f, delimiter=delim)
        for row in reader:
            rows.append(" | ".join(row))
    return [ParsedSection(text=_normalize("\n".join(rows)))]


# ----- HTML ------------------------------------------------------------------

@register(".html", ".htm", ".xhtml")
def parse_html(path: Path) -> list[ParsedSection]:
    try:
        from bs4 import BeautifulSoup  # type: ignore
    except ImportError as e:
        raise RuntimeError("beautifulsoup4 not installed; required for HTML") from e
    raw = path.read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(raw, "lxml" if _has_lxml() else "html.parser")
    for t in soup(["script", "style", "noscript", "header", "footer", "nav"]):
        t.decompose()
    text = soup.get_text("\n", strip=True)
    return [ParsedSection(text=_normalize(text))]


def _has_lxml() -> bool:
    try:
        import lxml  # noqa: F401
        return True
    except ImportError:
        return False


# ----- EPUB (handled by stdlib zipfile) --------------------------------------

@register(".epub")
def parse_epub(path: Path) -> list[ParsedSection]:
    sections: list[ParsedSection] = []
    try:
        from bs4 import BeautifulSoup  # type: ignore
    except ImportError:
        BeautifulSoup = None  # type: ignore
    with zipfile.ZipFile(path) as z:
        names = [n for n in z.namelist() if n.lower().endswith((".xhtml", ".html", ".htm"))]
        names.sort()
        for name in names:
            try:
                raw = z.read(name).decode("utf-8", errors="replace")
            except Exception:
                continue
            if BeautifulSoup is not None:
                soup = BeautifulSoup(raw, "lxml" if _has_lxml() else "html.parser")
                for t in soup(["script", "style", "nav"]):
                    t.decompose()
                text = soup.get_text("\n", strip=True)
            else:
                text = re.sub(r"<[^>]+>", " ", raw)
            text = _normalize(text)
            if text:
                sections.append(ParsedSection(text=text, section=name))
    return sections


# ----- EML / MBOX ------------------------------------------------------------

@register(".eml")
def parse_eml(path: Path) -> list[ParsedSection]:
    msg = email.message_from_bytes(path.read_bytes())
    parts: list[str] = []
    subj = msg.get("subject", "")
    frm = msg.get("from", "")
    to = msg.get("to", "")
    parts.append(f"Subject: {subj}\nFrom: {frm}\nTo: {to}")
    for part in msg.walk():
        ctype = part.get_content_type()
        if ctype == "text/plain":
            payload = part.get_payload(decode=True)
            if payload:
                parts.append(payload.decode(part.get_content_charset() or "utf-8",
                                            errors="replace"))
        elif ctype == "text/html":
            payload = part.get_payload(decode=True)
            if payload and not any(p for p in parts[1:]):
                try:
                    from bs4 import BeautifulSoup  # type: ignore
                    soup = BeautifulSoup(payload, "lxml" if _has_lxml() else "html.parser")
                    parts.append(soup.get_text("\n", strip=True))
                except ImportError:
                    parts.append(re.sub(r"<[^>]+>", " ",
                                        payload.decode("utf-8", errors="replace")))
    return [ParsedSection(text=_normalize("\n\n".join(parts)))]


# ----- Images (delegate to OCR module) --------------------------------------

@register(".png", ".jpg", ".jpeg", ".webp", ".tiff", ".tif", ".bmp")
def parse_image(path: Path) -> list[ParsedSection]:
    from .ocr import ocr_file
    text = ocr_file(path)
    if not text.strip():
        return []
    return [ParsedSection(text=_normalize(text), meta={"ocr": True})]
